#!/usr/bin/env python3
"""
Fill a Harvard SEAS offer letter template (.docx) with recruit details.

INTERNATIONAL paragraph rule:
  - International visitor: keep the paragraph, strip "INTERNATIONAL:" prefix
  - Domestic visitor: delete the paragraph entirely

The INTERNATIONAL: prefix is split across Word runs as:
  run[0] = "INTERNATIONAL"  run[1] = ": "  run[2+] = body text
"""

import argparse
import os
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def is_international_para(para):
    return para.text.strip().startswith("INTERNATIONAL")


def strip_international_prefix(para):
    """Clear run[0] ('INTERNATIONAL') and run[1] (':') leaving body text intact."""
    runs = para.runs
    if len(runs) >= 2 and runs[0].text.strip() == "INTERNATIONAL":
        runs[0].text = ""
        runs[1].text = ""


def fill(src, out, name, email, topic, pi_name, pi_title,
         letter_date, start_date, end_date, is_international, phd_received=True):

    doc = Document(src)
    date_count = [0]
    paras_to_remove = []

    for para in doc.paragraphs:
        # Delete "Dr. FirstName LastName" header line — email line is kept
        if para.text.strip() in ("Dr. FirstName LastName", "Dr. FirstName LastName"):
            paras_to_remove.append(para)
            continue
        # Catch any variant where runs spell out Dr./FirstName/LastName
        runs_text = [r.text for r in para.runs]
        if runs_text and runs_text[0].strip() == "Dr." and any("FirstName" in r for r in runs_text):
            paras_to_remove.append(para)
            continue

        # Handle INTERNATIONAL paragraphs
        if is_international_para(para):
            if is_international:
                strip_international_prefix(para)
            else:
                paras_to_remove.append(para)
            continue

        # DATE substitution: positional (1=letter date, 2=start, 3=end)
        if "DATE" in para.text:
            for run in para.runs:
                while "DATE" in run.text:
                    date_count[0] += 1
                    if date_count[0] == 1:
                        replacement = letter_date
                    elif date_count[0] == 2:
                        replacement = start_date
                    elif date_count[0] == 3:
                        replacement = end_date
                    else:
                        break
                    run.text = run.text.replace("DATE", replacement, 1)

        # Dear NAME salutation — NAME is always its own run
        if para.text.strip().startswith("Dear ") and "NAME" in para.text:
            for run in para.runs:
                # Strip "Dr. " from salutation if recruit has not yet received PhD
                if not phd_received:
                    run.text = run.text.replace("Dear Dr. ", "Dear ")
                if run.text == "NAME":
                    run.text = name
            continue

        # Email address line — placeholder is split across two runs ("Email" + " Address")
        if para.text.strip() in ("Email Address", "EMAIL ADDRESS"):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = email
            continue

        # Standard field substitutions
        for run in para.runs:
            run.text = run.text.replace("Email Address", email)
            run.text = run.text.replace("EMAIL ADDRESS", email)
            run.text = run.text.replace("[TOPIC(S) OF RESEARCH]", topic)
            run.text = run.text.replace("signatureFaculty Name", pi_name)
            run.text = run.text.replace("Faculty Name", pi_name)
            run.text = run.text.replace("Title", pi_title)

    # Delete domestic INTERNATIONAL paragraphs
    for para in paras_to_remove:
        para._element.getparent().remove(para._element)

    # Fix signature block: left-align and remove all indents (left + firstLine)
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    in_signature = False
    for para in doc.paragraphs:
        if para.text.strip().lower().startswith("sincerely"):
            in_signature = True
        if in_signature:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = 0
            para.paragraph_format.first_line_indent = 0
            # Also zero out firstLine directly in XML in case python-docx doesn't clear it
            pPr = para._p.find(f"{{{W}}}pPr")
            if pPr is not None:
                ind = pPr.find(f"{{{W}}}ind")
                if ind is not None:
                    ind.attrib.pop(f"{{{W}}}firstLine", None)
                    ind.attrib[f"{{{W}}}left"] = "0"

    out = os.path.expanduser(out)
    doc.save(out)
    print(f"Saved: {out}")
    return out


def main():
    parser = argparse.ArgumentParser(description="Fill a SEAS offer letter template")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--name", required=True)
    parser.add_argument("--email", required=True)
    parser.add_argument("--topic", required=True)
    parser.add_argument("--pi-name", required=True)
    parser.add_argument("--pi-title", required=True)
    parser.add_argument("--letter-date", required=True)
    parser.add_argument("--start-date", required=True)
    parser.add_argument("--end-date", required=True)
    parser.add_argument("--international", required=True, choices=["true", "false"])
    parser.add_argument("--phd-received", default="true", choices=["true", "false"],
                        help="true = keep 'Dear Dr. NAME'; false = use 'Dear NAME'")
    args = parser.parse_args()

    fill(
        src=args.input,
        out=args.output,
        name=args.name,
        email=args.email,
        topic=args.topic,
        pi_name=args.pi_name,
        pi_title=args.pi_title,
        letter_date=args.letter_date,
        start_date=args.start_date,
        end_date=args.end_date,
        is_international=(args.international == "true"),
        phd_received=(args.phd_received == "true"),
    )


if __name__ == "__main__":
    main()
